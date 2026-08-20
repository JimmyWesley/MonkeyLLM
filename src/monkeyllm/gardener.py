# SPDX-License-Identifier: Apache-2.0
# Copyright 2026 Jimmy Wesley

"""The Gardener (spec v0.9, Part G): brownfield ingest — adopt + sync.

Four stages, only stage 2 ever needs an LLM (and v1 runs without one):

    0 archive  ->  1 convert  ->  2 curate  ->  3 plant

Trusted infrastructure: it writes through the same audited mechanics as
everything else (nodes via C.7 plant, datasets via C.7.1, sync updates via
a `.md`-only git commit). Converters are pluggable (G.2): forest-config
command hooks > `monkeyllm.converters` entry points > built-ins.
"""

from __future__ import annotations

import csv
import datetime as dt
import hashlib
import json
import os
import re
import shlex
import shutil
import sqlite3
import subprocess
import tempfile
import unicodedata
from dataclasses import dataclass, field
from importlib.metadata import entry_points
from pathlib import Path
from typing import Callable, Protocol

import yaml

from monkeyllm import indexer
from monkeyllm.errors import E_SCHEMA, VineError
from monkeyllm.models import (
    MANUAL_SECTION, SAMPLE_ROWS, SAMPLE_SECTION,
    dataset_map, rows_label, validate_summary,
)
from monkeyllm.parser import (
    append_section, extract_section, replace_section, serialize_node,
)
from monkeyllm.tokens import estimate_tokens
from monkeyllm.vine import Vine

GARDENER_CONFIG = "gardener.yaml"  # lives in _meta/ (not a node: non-.md)
FOREST_MARKER = "_index.md"  # A.5: what makes a directory a forest root
DEFAULT_IGNORES = (".git", ".svn", ".hg", "__pycache__", "node_modules",
                   "_derived", "_assets")
DEFAULT_IGNORE_GLOBS = ("~$*", "*.tmp", "*.lock", ".DS_Store", "Thumbs.db")
ASSETS_DIR = "_assets"
SUMMARY_TARGET_TOKENS = 50
INGEST_CONFIDENCE = 0.7  # G.4: unreviewed by an LLM or a human

# G.10.1: the phases a document passes through inside its one step. Closed
# and ordered on purpose — a consumer renders position as index/len, which
# is the only reason a one-document batch can show progress at all.
STAGE_CONVERT = "convert"
STAGE_CURATE = "curate"
STAGE_PLANT = "plant"
STAGES = (STAGE_CONVERT, STAGE_CURATE, STAGE_PLANT)

PAYLOAD_TYPE_BY_EXT = {
    ".pdf": "pdf", ".docx": "docx",
    ".png": "image", ".jpg": "image", ".jpeg": "image", ".gif": "image",
    ".webp": "image",
    ".mp3": "audio", ".wav": "audio", ".m4a": "audio", ".ogg": "audio",
    ".flac": "audio",
}

# G.5.1: the extensions the media stub claims — exactly the image and audio
# halves of PAYLOAD_TYPE_BY_EXT, kept as named sets because the typing rule
# and the staging-archive rule test the same membership.
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp"}
AUDIO_EXTENSIONS = {".mp3", ".wav", ".m4a", ".ogg", ".flac"}


# ===========================================================================
# G.2 — the converter contract (public plugin API v1)
# ===========================================================================

@dataclass
class Conversion:
    """What a converter hands back: markdown, a dataset description, or —
    for a format the forest already speaks — the payload itself (G.2.2).

    `tables`/`samples`/`counts` describe a `payload` conversion for the
    G.2.3 map: the structure read from the source, three rows per table,
    and the row counts. They are never the data — a payload conversion
    reads the shape of a database, never the whole of it.
    """

    kind: str  # "markdown" | "dataset" | "payload"
    title: str
    markdown: str = ""
    schema: dict | None = None          # C.7.1 declarative schema
    rows: dict[str, list[list]] | None = None
    tables: dict[str, dict[str, str]] | None = None
    samples: dict[str, list[list]] | None = None
    counts: dict[str, int] | None = None


class Converter(Protocol):
    extensions: set[str]

    def convert(self, path: Path) -> Conversion: ...


class MarkdownConverter:
    """Built-in: .md/.txt pass through — the body IS the content."""

    extensions = {".md", ".markdown", ".txt"}

    def convert(self, path: Path) -> Conversion:
        text = path.read_text(encoding="utf-8", errors="replace")
        m = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
        title = m.group(1).strip() if m else path.stem.replace("_", " ").replace("-", " ")
        return Conversion(kind="markdown", title=title, markdown=text)


def _infer_column_type(values: list[str]) -> str:
    """INTEGER < REAL < TEXT, judged over the sampled non-empty values.

    A native `float` is never INTEGER, whatever its fractional part: a
    spreadsheet hands over `99.5` as a number, and `int(99.5)` truncates
    silently where `int("99.5")` raises — so a column typed from strings
    and the same column typed from a workbook would disagree, and the
    workbook's version would lose money by rounding it.
    """
    kind = "INTEGER"
    seen = False
    for v in values:
        if v is None or v == "":
            continue
        seen = True
        if isinstance(v, (bool, int)):
            continue
        if isinstance(v, float):
            kind = "REAL" if kind != "TEXT" else kind
            continue
        try:
            int(str(v))
            continue
        except (TypeError, ValueError):
            pass
        try:
            float(str(v))
            kind = "REAL" if kind != "TEXT" else kind
        except (TypeError, ValueError):
            return "TEXT"
    return kind if seen else "TEXT"


def _coerce(value, sql_type: str):
    if value is None or value == "":
        return None
    try:
        if sql_type == "INTEGER":
            return int(value)
        if sql_type == "REAL":
            return float(value)
    except (TypeError, ValueError):
        pass
    return str(value)


def slugify(text: str) -> str:
    """Deterministic id segment: lowercase, ASCII-folded, [a-z0-9._-]."""
    folded = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode()
    slug = re.sub(r"[^a-z0-9._-]+", "-", folded.lower()).strip("-.")
    return slug or "node"


def _column_names(header: list[str]) -> list[str]:
    cols: list[str] = []
    used: set[str] = set()
    for i, name in enumerate(header):
        col = slugify(str(name) or f"col{i}").replace(".", "_").replace("-", "_")[:48]
        if not re.match(r"^[a-z_]", col):
            col = f"c_{col}"
        while col in used:
            col += "_"
        used.add(col)
        cols.append(col)
    return cols


def _tables_conversion(title: str,
                       tables: list[tuple[str, list[str], list[list]]]) -> Conversion:
    """One `dataset` conversion out of one or more (name, header, records)
    tables — a workbook's sheets (G.2.4) or a single delimited file."""
    schema: dict = {}
    rows: dict[str, list[list]] = {}
    for table, header, records in tables:
        cols = _column_names(header)
        types = {
            col: _infer_column_type([r[i] if i < len(r) else None for r in records])
            for i, col in enumerate(cols)
        }
        schema[table] = {"columns": types}
        rows[table] = [
            [_coerce(r[i] if i < len(r) else None, types[col])
             for i, col in enumerate(cols)]
            for r in records
        ]
    return Conversion(kind="dataset", title=title, schema=schema, rows=rows)


def _tabular_conversion(title: str, table: str, header: list[str],
                        records: list[list]) -> Conversion:
    return _tables_conversion(title, [(table, header, records)])


def _table_name(path: Path) -> str:
    return _sql_name(path.stem)


def _sql_name(text: str) -> str:
    """A table name out of arbitrary text — a filename, a sheet's tab."""
    name = slugify(text).replace(".", "_").replace("-", "_")[:48]
    return name if re.match(r"^[a-z_]", name) else f"t_{name}"


def _sheet_tables(sheets: list[tuple[str, list[list]]], path: Path,
                  ) -> list[tuple[str, list[str], list[list]]]:
    """G.2.4: every sheet becomes a table. Empty sheets are skipped.

    No count limit here (G.2.5): C.7.1's ≤10 tables and ≤50 columns bound
    what a MODEL declares, and a workbook somebody exported is not a
    declaration. Taking sheet one and dropping the rest — or refusing a
    141-column ERP export — is the tool telling the operator their data is
    wrong. The map's own caps (G.2.3) are what keep the body bounded.
    """
    tables: list[tuple[str, list[str], list[list]]] = []
    used: set[str] = set()
    for sheet_name, data in sheets:
        data = [row for row in data if any(v not in (None, "") for v in row)]
        if len(data) < 2:
            continue
        table = _sql_name(sheet_name) or _table_name(path)
        while table in used:
            table += "_"
        used.add(table)
        header = [str(v) if v not in (None, "") else f"col{i}"
                  for i, v in enumerate(data[0])]
        tables.append((table, header, data[1:]))
    if not tables:
        raise VineError(
            E_SCHEMA, f"workbook has no data rows: {path.name}",
            hint="Every sheet was empty or held only a header row.")
    return tables


class CsvConverter:
    """Built-in: .csv -> dataset (C.7.1 birth with inferred column types)."""

    extensions = {".csv"}

    def convert(self, path: Path) -> Conversion:
        text = path.read_text(encoding="utf-8-sig", errors="replace")
        try:
            dialect = csv.Sniffer().sniff(text[:4096], delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel
        reader = csv.reader(text.splitlines(), dialect)
        data = [row for row in reader if row]
        if len(data) < 2:
            raise VineError(E_SCHEMA, f"csv has no data rows: {path.name}")
        return _tabular_conversion(path.stem.replace("-", " ").replace("_", " "),
                                   _table_name(path), data[0], data[1:])


class JsonConverter:
    """Built-in: tabular .json (list of flat dicts) -> dataset;
    anything else -> markdown with the JSON embedded."""

    extensions = {".json"}

    def convert(self, path: Path) -> Conversion:
        title = path.stem.replace("-", " ").replace("_", " ")
        data = json.loads(path.read_text(encoding="utf-8-sig", errors="replace"))
        if (isinstance(data, list) and data
                and all(isinstance(r, dict) for r in data)
                and all(not isinstance(v, (dict, list)) for r in data for v in r.values())):
            header = list(dict.fromkeys(k for r in data for k in r))
            records = [[r.get(k) for k in header] for r in data]
            return _tabular_conversion(title, _table_name(path), header, records)
        body = f"# {title}\n\n```json\n{json.dumps(data, ensure_ascii=False, indent=2)}\n```\n"
        return Conversion(kind="markdown", title=title, markdown=body)


class XlsxConverter:
    """Built-in when openpyxl is importable: one table per sheet (G.2.4)."""

    extensions = {".xlsx"}

    def convert(self, path: Path) -> Conversion:
        from openpyxl import load_workbook  # optional dependency

        wb = load_workbook(path, read_only=True, data_only=True)
        try:
            sheets = []
            for ws in wb.worksheets:
                # A read-only worksheet trusts the file's own `<dimension>`
                # record, and files written by anything other than Excel
                # routinely declare `A1:A1` or omit it — openpyxl then
                # yields ONE row of a hundred and the converter reports a
                # workbook with no data. `reset_dimensions` makes it infer
                # the extent from the rows that are actually there, which
                # is the only source that cannot be wrong.
                if hasattr(ws, "reset_dimensions"):
                    ws.reset_dimensions()
                sheets.append(
                    (ws.title, [list(row) for row in ws.iter_rows(values_only=True)]))
        finally:
            wb.close()
        return _tables_conversion(path.stem.replace("-", " ").replace("_", " "),
                                  _sheet_tables(sheets, path))


class XlsConverter:
    """Built-in when xlrd is importable (G.2.4; xlrd is BSD-3, optional
    `ingest` extra — same gating as the openpyxl and python-docx built-ins).
    xlrd 2.x reads the legacy `.xls` format and only that, which is exactly
    the gap openpyxl leaves."""

    extensions = {".xls"}

    def convert(self, path: Path) -> Conversion:
        import xlrd  # optional dependency (ingest extra)

        book = xlrd.open_workbook(str(path))
        try:
            sheets = [
                (sheet.name,
                 [[sheet.cell_value(r, c) for c in range(sheet.ncols)]
                  for r in range(sheet.nrows)])
                for sheet in book.sheets()
            ]
        finally:
            if hasattr(book, "release_resources"):
                book.release_resources()
        return _tables_conversion(path.stem.replace("-", " ").replace("_", " "),
                                  _sheet_tables(sheets, path))


SQLITE_MAGIC = b"SQLite format 3\x00"


class SqliteConverter:
    """Built-in (G.2.2): a SQLite file IS a dataset payload, so it is
    adopted rather than converted.

    Nothing here reads the data: the structure of every table and its first
    `SAMPLE_ROWS` rows are what the G.2.3 map needs, and the largest thing
    this holds is `3 x columns` values per table. The bytes themselves are
    the Gardener's to install — copying a database is O(bytes), while
    rebuilding it row by row is unbounded in the source's size and lossy
    wherever its declared types, views, indexes or BLOBs do not survive a
    TEXT|INTEGER|REAL|BLOB round trip.
    """

    extensions = {".db", ".sqlite", ".sqlite3"}

    def convert(self, path: Path) -> Conversion:
        with path.open("rb") as fh:
            if fh.read(len(SQLITE_MAGIC)) != SQLITE_MAGIC:
                raise VineError(
                    E_SCHEMA, f"not a SQLite database: {path.name}",
                    hint="The file's header is not 'SQLite format 3'. An "
                         "encrypted or truncated database reads the same way.")
        tables, samples, counts = read_sqlite_map(path)
        if not tables:
            raise VineError(E_SCHEMA, f"database has no tables: {path.name}")
        return Conversion(
            kind="payload", title=path.stem.replace("-", " ").replace("_", " "),
            tables=tables, samples=samples, counts=counts)


def read_sqlite_map(db: Path) -> tuple[dict, dict, dict]:
    """The G.2.3 map of a SQLite file: structure, first rows, row counts.

    Tables only, and in name order: `look`'s `query_manual` (C.2) reads the
    payload the same way, and a body claiming a view the digest never
    mentions is two answers to "what is in this dataset". Name order also
    keeps the map stable, so a `sync` rewrites it only when the data moved.

    Read-only (`mode=ro`), and every per-table read is guarded on its own:
    one unreadable table (a virtual table whose module is not loaded here)
    costs its own row, never the whole map.
    """
    conn = sqlite3.connect(f"file:{db.as_posix()}?mode=ro", uri=True)
    conn.text_factory = bytes_or_str
    tables: dict[str, dict[str, str]] = {}
    samples: dict[str, list[list]] = {}
    counts: dict[str, int] = {}
    try:
        names = [r[0] for r in conn.execute(
            "SELECT name FROM sqlite_master WHERE type = 'table' "
            "AND name NOT LIKE 'sqlite_%' ORDER BY name")]
        for name in names:
            quoted = '"' + str(name).replace('"', '""') + '"'
            try:
                info = list(conn.execute(f"PRAGMA table_info({quoted})"))
            except sqlite3.Error:
                continue
            if not info:
                continue
            tables[str(name)] = {str(c[1]): str(c[2] or "") for c in info}
            try:
                cur = conn.execute(f"SELECT * FROM {quoted} LIMIT {SAMPLE_ROWS}")
                samples[str(name)] = [list(r) for r in cur.fetchall()]
                counts[str(name)] = conn.execute(
                    f"SELECT COUNT(*) FROM {quoted}").fetchone()[0]
            except sqlite3.Error:
                samples[str(name)] = []
    finally:
        conn.close()
    return tables, samples, counts


def bytes_or_str(raw: bytes):
    """SQLite's text factory for a foreign database: decode when it is text,
    keep the bytes when it is not. A source nobody in this project created
    may hold any encoding, and a UnicodeDecodeError inside the map would
    lose a table that reads perfectly well through `query`."""
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw


class DocxConverter:
    """Built-in when python-docx is importable (G.2.1, MIT/BSD-clean):
    single-pass w:t traversal in document order — heading-styled paragraphs,
    pipe tables, and text inside embedded text boxes (wps:txbx/v:textbox);
    headers/footers excluded (letterhead boilerplate is scent noise).
    Technique derived from the owner's pdf-replace reader."""

    extensions = {".docx"}

    def convert(self, path: Path) -> Conversion:
        from docx import Document  # optional dependency (ingest extra)
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph

        doc = Document(str(path))
        tag_p, tag_tbl = qn("w:p"), qn("w:tbl")
        tag_t, tag_tr, tag_tc = qn("w:t"), qn("w:tr"), qn("w:tc")

        def text_of(el) -> str:
            # joining EVERY descendant w:t merges runs Word fragmented
            # mid-word and captures text living inside embedded text boxes
            joined = "".join(t.text or "" for t in el.iter(tag_t))
            return re.sub(r"\s+", " ", joined).strip()

        title = ""
        lines: list[str] = []
        for block in doc.element.body.iterchildren():
            if block.tag == tag_p:
                text = text_of(block)
                if not text:
                    continue
                level = self._heading_level(Paragraph(block, doc))
                if level == 1 and not title:
                    title = text
                # node title owns "#"; document headings start at "##"
                lines.append(f"{'#' * min(level + 1, 6)} {text}" if level else text)
                lines.append("")
            elif block.tag == tag_tbl:
                # direct children only: nested tables flatten into cell text
                rows = [
                    [text_of(tc).replace("|", "\\|") for tc in tr.findall(tag_tc)]
                    for tr in block.findall(tag_tr)
                ]
                rows = [r for r in rows if any(r)]
                if not rows:
                    continue
                width = max(len(r) for r in rows)
                rows = [r + [""] * (width - len(r)) for r in rows]
                lines.append("| " + " | ".join(rows[0]) + " |")
                lines.append("|" + " --- |" * width)
                lines.extend("| " + " | ".join(r) + " |" for r in rows[1:])
                lines.append("")
        title = title or path.stem.replace("_", " ").replace("-", " ")
        content = "\n".join(lines).strip()
        body = f"# {title}\n\n{content}\n" if content else f"# {title}\n"
        return Conversion(kind="markdown", title=title, markdown=body)

    @staticmethod
    def _heading_level(para) -> int:
        name = (para.style.name if para.style is not None else "") or ""
        if name == "Title":
            return 1
        m = re.match(r"Heading (\d)$", name)
        return int(m.group(1)) if m else 0


class CommandConverter:
    """G.2 discovery source 1: an external command template from the forest
    config converts the file — any tool, any license, never our dependency."""

    def __init__(self, extension: str, template: str):
        self.extensions = {extension.lower()}
        self.template = template

    def convert(self, path: Path) -> Conversion:
        with tempfile.TemporaryDirectory(prefix="gardener-") as tmp:
            out = Path(tmp) / (path.stem + ".md")
            # non-posix split keeps Windows backslashes; strip the quotes it
            # leaves around tokens, then substitute placeholders post-split
            parts = [
                p[1:-1] if len(p) >= 2 and p[0] == p[-1] and p[0] in "\"'" else p
                for p in shlex.split(self.template, posix=False)
            ]
            cmd = [part.replace("{input}", str(path)).replace("{output}", str(out))
                   for part in parts]
            r = subprocess.run(cmd, capture_output=True, text=True, timeout=600)
            if r.returncode != 0:
                raise VineError(
                    E_SCHEMA,
                    f"converter command failed for {path.name} (exit {r.returncode})",
                    hint=(r.stderr or r.stdout or "").strip()[:200],
                )
            text = out.read_text(encoding="utf-8", errors="replace") if out.is_file() else r.stdout
            if not text.strip():
                raise VineError(E_SCHEMA, f"converter command produced no markdown: {path.name}")
        m = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
        title = m.group(1).strip() if m else path.stem.replace("_", " ").replace("-", " ")
        return Conversion(kind="markdown", title=title, markdown=text)


# G.5.1 / H.3 (v0.54): the stub's admission, shared with the Ranger's
# `needs_description` check — two spellings of a sentinel agree only where
# somebody compared them.
MEDIA_STUB_SENTINEL = "No description has been generated for this media yet."


class MediaStubConverter:
    """Built-in (G.5.1): the model-free floor for image and audio files.

    Before this existed an image was `unsupported` — no converter claimed
    it, so a screenshot fell out of the report entirely. The stub returns
    the only markdown that needs no model: what the file is called, what
    format it is, how big it is, and the plain admission that nothing has
    described it yet. A richer converter injected ahead of this one (the
    `extra_converters` seam) replaces the body; the stub is what guarantees
    the node exists either way.
    """

    extensions = IMAGE_EXTENSIONS | AUDIO_EXTENSIONS

    def convert(self, path: Path) -> Conversion:
        title = path.stem.replace("_", " ").replace("-", " ")
        size = path.stat().st_size
        fmt = path.suffix.lstrip(".").upper()
        body = (
            f"# {title}\n\n"
            f"Media file `{path.name}` ({fmt} format, {size} bytes).\n\n"
            f"{MEDIA_STUB_SENTINEL}\n"
        )
        return Conversion(kind="markdown", title=title, markdown=body)


def builtin_converters() -> list:
    # SQLite needs no optional dependency: it is the standard library and it
    # is already this project's payload format (G.2.2).
    convs: list = [MarkdownConverter(), CsvConverter(), JsonConverter(),
                   SqliteConverter()]
    try:
        import openpyxl  # noqa: F401

        convs.append(XlsxConverter())
    except ImportError:
        pass
    try:
        import xlrd  # noqa: F401

        convs.append(XlsConverter())
    except ImportError:
        pass
    try:
        import docx  # noqa: F401

        convs.append(DocxConverter())
    except ImportError:
        pass
    # G.5.1: last on purpose — the stub is the floor every richer media
    # converter (a command hook, an injected describer) stands above.
    convs.append(MediaStubConverter())
    return convs


def discover_converters(config: dict, extra: list | None = None) -> list:
    """G.2 order: config command hooks > injected extras (G.5.1) > entry
    points > built-ins.

    `extra` is the seam a host uses to inject converters it holds (the
    vision describer): AFTER the operator's command hooks — an operator who
    configured their own `.png` hook keeps it — and BEFORE entry points and
    built-ins, so everyone else gets the injected converter over the stub.
    """
    convs: list = [
        CommandConverter(ext, tpl)
        for ext, tpl in (config.get("converters") or {}).items()
    ]
    convs.extend(extra or [])
    for ep in entry_points(group="monkeyllm.converters"):
        try:
            loaded = ep.load()
            convs.append(loaded() if isinstance(loaded, type) else loaded)
        except Exception:  # a broken plugin never blocks the pipeline
            continue
    convs.extend(builtin_converters())
    return convs


def discover_hooks() -> list[Callable]:
    """G.4.3: `on_curate` hooks from the `monkeyllm.hooks` entry-point group."""
    hooks: list[Callable] = []
    for ep in entry_points(group="monkeyllm.hooks"):
        if ep.name != "on_curate":
            continue
        try:
            hooks.append(ep.load())
        except Exception:
            continue
    return hooks


# ===========================================================================
# G.4.1 — LLM-free curation
# ===========================================================================

_MD_NOISE = re.compile(r"^#+\s+.*$|^[-*>|`].*$|!\[[^\]]*\]\([^)]*\)", re.MULTILINE)
_MD_INLINE = re.compile(r"\[([^\]]*)\]\([^)]*\)|[*_`]{1,3}")


def derive_aliases(rel: Path, alias_map: dict) -> list[str]:
    """G.2.6 (v0.54): the team's own name for a document.

    Mechanical and declared, never guessed: a source file whose stem starts
    with digits, sitting in a folder the operator's `aliases:` map names,
    gains the conversational form (`BE-291`) and the path form
    (`back-end/291`) — the two spellings integrators were observed to try.
    No map, no aliases: the convention is content vocabulary, and content
    vocabulary lives in the forest's own config, never in the engine.
    """
    if not isinstance(alias_map, dict) or not alias_map:
        return []
    m = re.match(r"(\d+)", rel.stem)
    if m is None:
        return []
    folder = rel.parent.name
    prefix = alias_map.get(folder)
    if not prefix or not isinstance(prefix, str):
        return []
    num = m.group(1)
    return [f"{prefix}-{num}", f"{folder}/{num}"]


def derive_summary(markdown: str, title: str) -> str:
    """First meaningful sentences of the content, <= 60 tokens (A.4)."""
    text = _MD_INLINE.sub(r"\1", _MD_NOISE.sub("", markdown))
    text = re.sub(r"\s+", " ", text).strip()
    summary = ""
    for sentence in re.split(r"(?<=[.!?])\s+", text):
        candidate = f"{summary} {sentence}".strip()
        if summary and estimate_tokens(candidate) > SUMMARY_TARGET_TOKENS:
            break
        summary = candidate
        if estimate_tokens(summary) > SUMMARY_TARGET_TOKENS:
            words = summary.split()
            while words and estimate_tokens(" ".join(words)) > SUMMARY_TARGET_TOKENS:
                words.pop()
            summary = " ".join(words).rstrip(",;") + "…"
            break
    summary = summary or f"Adopted content '{title}'; pending curation."
    try:
        validate_summary(summary)
    except VineError:
        summary = f"Adopted content '{title}'; pending curation."
    return summary


def derive_branch_summary(title: str, child_titles: list[str]) -> str:
    """Deterministic G.4.4 fallback: compose the region's scent from child
    titles, <= 60 tokens (A.4). Never raises."""
    names = [t.strip() for t in child_titles if t and t.strip()]
    summary = f"Region '{title}' with {len(names)} entries."
    listing: list[str] = []
    for name in names:
        candidate = (f"Region '{title}': " + ", ".join(listing + [name])
                     + f" (+{len(names) - len(listing) - 1} more).")
        if listing and estimate_tokens(candidate) > SUMMARY_TARGET_TOKENS:
            break
        listing.append(name)
    if listing:
        more = len(names) - len(listing)
        tail = f" (+{more} more)." if more > 0 else "."
        summary = f"Region '{title}': " + ", ".join(listing) + tail
    try:
        validate_summary(summary)
    except VineError:
        summary = f"Region '{title}' with {len(names)} entries."
    return summary


# ===========================================================================
# The Gardener
# ===========================================================================

@dataclass
class IngestReport:
    planted: list[str] = field(default_factory=list)
    branches: list[str] = field(default_factory=list)
    updated: list[str] = field(default_factory=list)
    unchanged: list[str] = field(default_factory=list)
    stale: list[str] = field(default_factory=list)
    unsupported: list[str] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)
    # Dry-run only (spec J.8.1): the passports a real run would have planted,
    # in the order it would have planted them. Empty on every ordinary run,
    # so `as_dict` keeps reporting the same shape it always did.
    drafts: list[dict] = field(default_factory=list)

    def as_dict(self) -> dict:
        return {k: list(v) for k, v in self.__dict__.items()}


# G.10: which report list a step grew names what the step did. Ordered by
# precedence — a planted file may also record its branch, and the branch is
# not the action. A draft counts as `planted`: a preview steps exactly as
# the run it previews.
_STEP_ACTIONS = (("planted", "planted"), ("updated", "updated"),
                 ("unchanged", "unchanged"), ("unsupported", "unsupported"),
                 ("errors", "error"), ("stale", "stale"), ("drafts", "planted"))


def _counts(report: IngestReport) -> dict:
    return {attr: len(getattr(report, attr)) for attr, _ in _STEP_ACTIONS}


def _step(file: str, index: int, total: int, report: IngestReport,
          before: dict) -> dict:
    action = "skipped"
    for attr, name in _STEP_ACTIONS:
        if len(getattr(report, attr)) > before[attr]:
            action = name
            break
    return {"file": file, "index": index, "total": total, "action": action}


def _drain(steps: "IngestSteps") -> dict:
    for _ in steps:
        pass
    return steps.result


class IngestSteps:
    """G.10 step iterator: one document per `next()`, the report at the end.

    Construction is eager where iteration is lazy — the source is resolved
    and walked before this exists, so `total` is known up front and a bad
    source fails before any step runs (spec J.9 needs both to refuse before
    accepting). `report` is the live `IngestReport` the steps are filling —
    what a consumer that died mid-batch can still account from — and
    `result` is its final dict once the iterator is exhausted, None until
    then.
    """

    def __init__(self, total: int, steps, report: IngestReport):
        self.total = total
        self.report = report
        self.result: dict | None = None
        self._steps = steps

    def __iter__(self) -> "IngestSteps":
        return self

    def __next__(self) -> dict:
        try:
            return next(self._steps)
        except StopIteration as done:
            self.result = done.value
            raise


class Gardener:
    """Adopts a directory into forest (Part G).

    `dry_run` makes the whole object incapable of writing: it converts,
    curates and proposes exactly as a real run does, and then collects the
    drafts instead of planting them (spec J.8.1). The flag lives here rather
    than on `adopt`/`sync` on purpose — a per-call flag is forgotten by the
    next call somebody adds, and the guarantee this exists to give ("nothing
    was written") has to hold for the object, not for one entry point.
    """

    def __init__(self, vine: Vine, converters: list | None = None,
                 hooks: list[Callable] | None = None, *, dry_run: bool = False,
                 on_stage: Callable[[str, str], None] | None = None,
                 extra_converters: list | None = None,
                 provenance: dict[str, str] | None = None):
        self.vine = vine
        self.forest = vine.forest
        self.config = self._load_config()
        # J.8 (v0.48): source path -> URL, for sources whose origin is an
        # address rather than a directory (a clipped page, a saved image).
        # A MAP handed at construction, deliberately not an `on_curate`
        # hook: curation never runs on refreshes (G.3), so provenance
        # recorded there would vanish with the first `sync` — the map is
        # consulted on adopt and on every body refresh alike. Keys are the
        # relative posix paths `source_path` records; the URL is data, not
        # vocabulary, so nothing here reads it.
        self.provenance = dict(provenance or {})
        # G.5.1 seam: `extra_converters` joins discovery between the
        # operator's command hooks and everything else. An explicit
        # `converters` list bypasses discovery entirely (tests do this),
        # so the extras are ignored there — the caller already said
        # exactly what runs.
        self.converters = (converters if converters is not None
                           else discover_converters(self.config,
                                                    extra=extra_converters))
        self.hooks = hooks if hooks is not None else discover_hooks()
        self.dry_run = bool(dry_run)
        self.on_stage = on_stage

    def _stage(self, file: str, stage: str) -> None:
        """G.10.1: name the phase, never pause in it.

        A step is still a whole document — nothing here is a suspension
        point — but a batch of ONE document would otherwise show a consumer
        nothing until it shows everything, which is indistinguishable from
        a hang. An observer that raises is swallowed: progress reporting
        that can abort the work it reports on is worse than none.
        """
        if self.on_stage is None:
            return
        try:
            self.on_stage(file, stage)
        except Exception:
            pass

    # -- config (G.6) -------------------------------------------------------

    def _config_path(self) -> Path:
        return self.forest.root / "_meta" / GARDENER_CONFIG

    def _load_config(self) -> dict:
        p = self._config_path()
        if p.is_file():
            return yaml.safe_load(p.read_text(encoding="utf-8")) or {}
        return {}

    def _save_config(self) -> None:
        if self.dry_run:
            return  # a preview that recorded a source root would misdirect
        p = self._config_path()
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(yaml.safe_dump(self.config, allow_unicode=True, sort_keys=False),
                     encoding="utf-8")

    # -- walking ------------------------------------------------------------

    def _ignored(self, path: Path) -> bool:
        if any(part in DEFAULT_IGNORES or part.startswith(".") for part in path.parts):
            return True
        globs = list(DEFAULT_IGNORE_GLOBS) + list(self.config.get("ignore") or [])
        return any(path.match(g) for g in globs)

    def _walk(self, src: Path) -> list[Path]:
        """Every ingestable file under `src`, forests excluded.

        A forest met inside the tree is pruned whole: its passports are
        somebody's curated nodes, not documents to convert, and a source
        that happens to sit above a registry would otherwise hand every
        forest under it to this one — across the tenant boundary, in a
        single call.
        """
        out: list[Path] = []
        for dirpath, dirnames, filenames in os.walk(src):
            here = Path(dirpath)
            rel_dir = here.relative_to(src)
            if here != src and (here / FOREST_MARKER).is_file():
                dirnames[:] = []
                continue
            # Pruning the directory beats filtering its files one by one:
            # os.walk does not descend into what it is not given.
            dirnames[:] = [d for d in dirnames if not self._ignored(rel_dir / d)]
            for name in filenames:
                if not self._ignored(rel_dir / name):
                    out.append(here / name)
        return sorted(out)

    def _resolve_source(self, source: str | Path | None,
                        *, recorded: bool = False) -> Path:
        """The one place a host path becomes an ingest root (G.3).

        Every caller passes through here, so the two ways a walk can escape
        its forest are closed once: an empty source (which `Path("")`
        resolves to the process's working directory — the Station's own
        install tree, or whatever a shell happened to be sitting in), and a
        source that contains the forest, which walks the registry beside it.
        """
        raw = str(source or "").strip()
        if not raw and recorded:
            raw = str(self.config.get("source_root") or "").strip()
        if not raw:
            # Never fall back to the working directory: the caller asked to
            # ingest "the usual place" and this forest has no usual place.
            raise VineError(
                E_SCHEMA, "this forest has no adopted source to sync",
                hint="Adopt a directory first, or pass the source explicitly.")
        src = Path(raw).resolve()
        if not src.is_dir():
            raise VineError(
                E_SCHEMA, f"source is not a directory: {src}",
                hint="Adopt a directory first, or pass the source explicitly.")
        root = self.forest.root.resolve()
        if root == src or root.is_relative_to(src):
            raise VineError(
                E_SCHEMA, f"source contains the forest itself: {src}",
                hint="Ingest reads a source tree into a forest; a source at "
                     "or above the forest root would ingest the forest, and "
                     "every other forest beside it.")
        if src.is_relative_to(root) and not src.is_relative_to(root / "_derived"):
            # `_derived/` is the exception on purpose: the Station's upload
            # staging lives there, and it is explicitly not forest content.
            raise VineError(
                E_SCHEMA, f"source is inside the forest: {src}",
                hint="A forest's own nodes are not a source to re-ingest.")
        return src

    def _converter_for(self, path: Path):
        ext = path.suffix.lower()
        for conv in self.converters:
            if ext in conv.extensions:
                return conv
        return None

    def _with_provenance(self, markdown: str, rel: str) -> str:
        """J.8 (v0.48): a converted body whose source has an address ends
        with the same `Source:` line a composed clip carries.

        Applied to markdown conversions only — a dataset's map is not
        prose, and a payload's body is the G.2.3 sample map — and BEFORE
        curation and the content policy, so the Curator reads what a
        reader will and a cached body carries its address too. Idempotent
        by the STAMPED LINE, not by substring: a body citing a deeper
        link on the same site (`…/blog/post-123` under source
        `…/blog`) contains the URL as a prefix, and a substring test
        would silently drop the provenance for exactly the common case.
        Only a body already carrying the exact `Source:` line is left
        alone.
        """
        url = self.provenance.get(rel)
        if not url:
            return markdown
        if re.search(rf"(?m)^Source: {re.escape(url)}[ \t]*$", markdown):
            return markdown
        return markdown.rstrip("\n") + f"\n\n---\n\nSource: {url}\n"

    # -- ids and branches ----------------------------------------------------

    def _branch_id_for(self, rel_dir: Path, dest: str | None) -> str:
        parts = [slugify(p) for p in rel_dir.parts]
        prefix = [] if not dest else [dest]
        return "/".join(prefix + parts + ["_index"]) if (parts or prefix) else "_index"

    def _ensure_branch(self, rel_dir: Path, dest: str | None,
                       report: IngestReport) -> str:
        branch_id = self._branch_id_for(rel_dir, dest)
        if branch_id == "_index" or self.forest.exists(branch_id):
            return branch_id
        parent_id = self._ensure_branch(rel_dir.parent, dest, report) \
            if rel_dir.parts else "_index"
        name = rel_dir.parts[-1] if rel_dir.parts else dest
        title = str(name).replace("_", " ").replace("-", " ")
        if self.dry_run:
            # `branches` becomes "would create" — the id is still returned so
            # the drafts below name the parent the real run would give them.
            report.branches.append(branch_id)
            return branch_id
        self.vine.plant({
            "id": branch_id,
            "type": "branch",
            "parent": parent_id,
            "title": title,
            "summary": f"Documents adopted from source folder '{name}'.",
            "source": "ingest",
            "body": (f"# {title}\n\n> Documents adopted from source folder "
                     f"'{name}'.\n\n## Sub-branches\n\n## Direct bananas\n\n"
                     "## Cross trails\n"),
        })
        report.branches.append(branch_id)
        return branch_id

    # -- rollup (G.4.4) ------------------------------------------------------

    def rollup(self, curator=None, *, only_ingest: bool = True) -> dict:
        """G.4.4: synthesize branch summaries bottom-up (deepest first) from
        the children's entry lines. Writes through C.8 graft, so parent-entry
        propagation and `.md`-only commits are inherited."""
        if self.dry_run:
            return {"rolled": [], "fallbacks": [], "skipped": 0}
        rolled: list[str] = []
        fallbacks: list[str] = []
        skipped = 0
        rows = self.vine.catalog.conn.execute(
            "SELECT id, source, title, summary FROM nodes WHERE kind = 'branch' "
            "ORDER BY LENGTH(id) - LENGTH(REPLACE(id, '/', '')) DESC, id"
        ).fetchall()
        for row in rows:
            branch_id = row["id"]
            if branch_id.startswith("_meta/"):
                continue
            if only_ingest and row["source"] != "ingest":
                skipped += 1
                continue
            # Fresh read: a deeper child's rollup may have just rewritten
            # this branch's entry lines via summary propagation.
            node = self.forest.read(branch_id)
            entries: list[str] = []
            for section in (indexer.SUBBRANCH_SECTION, indexer.BANANAS_SECTION):
                sec = extract_section(node.body, section) or ""
                entries += [l for l in sec.splitlines() if l.startswith("- [[")]
            if not entries:
                skipped += 1  # empty region: the template summary stays
                continue
            new_summary = (curator.branch_summary(row["title"], entries)
                           if curator is not None else None)
            if new_summary is None:
                child_titles = [c["title"]
                                for c in self.vine.catalog.children(branch_id)]
                new_summary = derive_branch_summary(row["title"], child_titles)
                fallbacks.append(branch_id)
            if new_summary == row["summary"]:
                continue
            self.vine.graft(branch_id,
                            {"set_frontmatter": {"summary": new_summary}})
            rolled.append(branch_id)
        return {"rolled": rolled, "fallbacks": fallbacks, "skipped": skipped}

    # -- curation (G.4) -----------------------------------------------------

    def _curate(self, draft: dict, report: IngestReport) -> dict:
        draft.setdefault("tags", [])
        for tag in (self.config.get("curation") or {}).get("default_tags") or []:
            if tag not in draft["tags"]:
                draft["tags"].append(tag)
        for hook in self.hooks:
            try:
                result = hook(draft)
                if isinstance(result, dict):
                    draft = result
            except Exception as e:  # G.4.3: a broken hook never aborts ingest
                report.errors.append(f"on_curate hook {getattr(hook, '__name__', hook)!r}: {e}")
        return draft

    # -- stage 0: archive ----------------------------------------------------

    def _archive(self, src_file: Path, branch_id: str) -> tuple[str, str | None, str]:
        """Copy the original under the branch's _assets/ (gitignored).
        Returns (payload, payload_type | None, payload_hash)."""
        branch_dir = self.forest.path_for(branch_id).parent
        assets = branch_dir / ASSETS_DIR
        data = src_file.read_bytes()
        digest = hashlib.sha256(data).hexdigest()
        name = f"{digest[:8]}-{slugify(src_file.stem)}{src_file.suffix.lower()}"
        # The digest and the name are computed either way, so a dry run's
        # draft carries the same payload fields a real one would write.
        if not self.dry_run:
            assets.mkdir(parents=True, exist_ok=True)
            (assets / name).write_bytes(data)
        ptype = PAYLOAD_TYPE_BY_EXT.get(src_file.suffix.lower())
        return f"{ASSETS_DIR}/{name}", ptype, digest

    # -- adopt (G.3) ---------------------------------------------------------

    def adopt(self, source: str | Path, dest: str | None = None) -> dict:
        return _drain(self.adopt_iter(source, dest))

    def adopt_iter(self, source: str | Path,
                   dest: str | None = None) -> IngestSteps:
        """One document per step (G.10); `adopt` is exactly "drain this".

        The source root is recorded before the first step, not after the
        last: a run abandoned at any yield — a crash, a cancel (J.9) —
        leaves the stepped files planted and committed, and the recorded
        root is what lets `sync` finish the remainder instead of the
        operator starting over.
        """
        src = self._resolve_source(source)
        if not self.dry_run:
            self.config["source_root"] = src.as_posix()
            if dest:
                self.config["dest"] = dest
            self._save_config()
        files = self._walk(src)
        report = IngestReport()

        def steps():
            for i, f in enumerate(files):
                before = _counts(report)
                self._ingest_file(src, f, dest, report)
                yield _step(f.relative_to(src).as_posix(), i + 1, len(files),
                            report, before)
            return report.as_dict()

        return IngestSteps(len(files), steps(), report)

    def _ingest_file(self, src: Path, f: Path, dest: str | None,
                     report: IngestReport) -> None:
        rel = f.relative_to(src)
        ext = f.suffix.lower()
        claimants = [c for c in self.converters if ext in c.extensions]
        if not claimants:
            report.unsupported.append(rel.as_posix())
            return
        self._stage(rel.as_posix(), STAGE_CONVERT)
        # G.5.1: a converter that fails falls THROUGH, not out — the next
        # claimant in discovery order gets the file (a describer whose
        # endpoint is down falls back to the stub), and every failure lands
        # in the report's errors, naming who failed on what. Only when the
        # LAST claimant fails does the file take the terminal error path,
        # with the message shape it always had (G.2: a broken converter
        # never crashes adopt).
        conversion = None
        for conv_obj in claimants:
            terminal = conv_obj is claimants[-1]
            try:
                conversion = conv_obj.convert(f)
                break
            except VineError as e:
                if terminal:
                    report.errors.append(f"{rel.as_posix()}: {e.message}")
                    return
                report.errors.append(
                    f"{rel.as_posix()}: {type(conv_obj).__name__} failed, "
                    f"falling back: {e.message}")
            except Exception as e:
                if terminal:
                    report.errors.append(f"{rel.as_posix()}: converter error: {e}")
                    return
                report.errors.append(
                    f"{rel.as_posix()}: {type(conv_obj).__name__} failed, "
                    f"falling back: {e}")

        branch_id = self._ensure_branch(rel.parent, dest, report)
        node_id = self._node_id(rel, dest)
        # A real run collides against nodes that now exist; a dry run has to
        # count the drafts too, or two previewed files that slug the same way
        # would both claim the id and only one of them would be right.
        taken = self.forest.exists(node_id) or (
            self.dry_run and any(d["id"] == node_id for d in report.drafts))
        if taken:
            node_id = f"{node_id}-{hashlib.sha256(rel.as_posix().encode()).hexdigest()[:6]}"

        st = f.stat()
        source_hash = hashlib.sha256(f.read_bytes()).hexdigest()
        is_text_source = ext in MarkdownConverter.extensions
        # G.5.1: media is typed off the SOURCE, not off what the converter
        # returned — a describer and the stub both hand back markdown, and
        # the passport has to say `media` either way.
        is_media_source = PAYLOAD_TYPE_BY_EXT.get(ext) in ("image", "audio")
        draft: dict = {
            "id": node_id,
            "parent": branch_id,
            "title": conversion.title,
            "source": "ingest",
            "confidence": INGEST_CONFIDENCE,
            "source_path": rel.as_posix(),
            "source_hash": source_hash,
            # G.8 fast-path: sync skips hashing when both still match
            "source_size": st.st_size,
            "source_mtime": round(st.st_mtime, 3),
        }
        # G.2.6 (v0.54): part of the draft build, so adopt derives them and
        # sync recomputes rather than erases. Curation never touches them.
        aliases = derive_aliases(rel, self.config.get("aliases") or {})
        if aliases:
            draft["aliases"] = aliases
        if conversion.kind == "dataset":
            # The map goes in here rather than being left to C.7.1's auto
            # manual, because curation runs BEFORE the plant and G.4.6 reads
            # it. Same text either way — plant keeps a caller-provided
            # manual verbatim — so the body a reader sees is unchanged.
            draft.update({
                "type": "dataset",
                "schema": conversion.schema,
                "rows": conversion.rows,
                "body": f"# {conversion.title}\n\n" + dataset_map(
                    {t: {c: ty.upper() for c, ty in s["columns"].items()}
                     for t, s in (conversion.schema or {}).items()},
                    conversion.rows,
                    {t: len(r) for t, r in (conversion.rows or {}).items()}),
                "summary": self._dataset_summary(conversion),
            })
        elif conversion.kind == "payload":
            # G.2.2: the source IS the payload. The body carries the map
            # here rather than letting C.7.1 generate it — plant is not
            # creating this database and has never read it.
            draft.update({
                "type": "dataset",
                "payload_type": "sqlite",
                "body": f"# {conversion.title}\n\n"
                        + dataset_map(conversion.tables or {}, conversion.samples,
                                      conversion.counts),
                "summary": self._dataset_summary(conversion),
            })
        else:
            # G.5.1 typing rule: text source -> note; an image/audio source
            # -> media; every other converted format -> document.
            draft.update({
                "type": ("note" if is_text_source
                         else "media" if is_media_source else "document"),
                "body": conversion.markdown,
                "summary": derive_summary(conversion.markdown, conversion.title),
            })
        # G.7 archive policy: durable sources are referenced, not copied.
        # A payload conversion's original IS its payload (G.2.2 rule 6) —
        # archiving it would store the same bytes twice under two hashes.
        # G.5.1 amendment: media STAGED under this forest's `_derived/` (an
        # upload) is archived regardless of the policy — `_derived/` is
        # disposable by contract, so the `_assets/` copy is the only one
        # that will exist. Resolved paths on both sides: the forest root
        # may sit behind a symlink (same comparison `_resolve_source` makes).
        staged_media = is_media_source and src.resolve().is_relative_to(
            self.forest.root.resolve() / "_derived")
        if (not is_text_source and conversion.kind != "payload"
                and (staged_media
                     or self.config.get("archive", "never") == "always")):
            payload, ptype, phash = self._archive(f, branch_id)
            # dataset payload is its own .db; unknown payload types are
            # archived but not referenced (the A.3 enum stays honest)
            if conversion.kind != "dataset" and ptype:
                draft.update({"payload": payload, "payload_type": ptype,
                              "payload_hash": phash})

        # J.8 (v0.48): the address is stamped after the summary derived —
        # a URL is not scent, and locate's 60 tokens are too few to spend
        # on one — but before curation and the content policy, so the
        # Curator reads what a reader will and a cached body carries its
        # address too.
        if conversion.kind == "markdown":
            draft["body"] = self._with_provenance(draft["body"], rel.as_posix())

        # curation sees the FULL converted text (G.7.4)…
        self._stage(rel.as_posix(), STAGE_CURATE)
        draft = self._curate(draft, report)
        # …and only then the content policy slims the node (G.7)
        if conversion.kind == "markdown":
            draft = self._apply_content_policy(draft, conversion.title,
                                               is_text_source)
        if self.dry_run:
            report.drafts.append(draft)
            return
        self._stage(rel.as_posix(), STAGE_PLANT)
        installed: Path | None = None
        try:
            if conversion.kind == "payload":
                installed = self._install_payload(node_id, f)
                draft["payload"] = installed.name
                draft["payload_hash"] = source_hash
            # G.2.5: the Gardener is trusted infrastructure adopting a
            # source that already exists, not a model declaring a schema —
            # so the C.7.1 count limits do not bind it. Names and types are
            # validated exactly as they are for everyone else.
            self.vine.plant(draft, adopted=True)
            report.planted.append(node_id)
        except VineError as e:
            # C.7's atomicity extends to the copy (G.2.2 rule 5): a payload
            # whose passport was refused is a file nothing references.
            if installed is not None:
                installed.unlink(missing_ok=True)
            report.errors.append(f"{rel.as_posix()}: {e.message}")
        except Exception:
            if installed is not None:
                installed.unlink(missing_ok=True)
            raise

    def _install_payload(self, node_id: str, source: Path) -> Path:
        """G.2.2 rule 2: the source database, copied beside its passport
        under the bare `<leaf>.db` name a C.7.1 birth would have used."""
        node_path = self.forest.path_for(node_id)
        db = node_path.parent / f"{node_path.stem}.db"
        if db.exists():
            raise VineError(
                E_SCHEMA, f"payload already exists: {db.name}",
                hint="An adopted database never overwrites a payload.")
        db.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(source, db)
        return db

    @staticmethod
    def _refresh_map(body: str, tables: dict, samples: dict | None,
                     counts: dict | None) -> str:
        """G.2.3 rule 4: rewrite the two generated sections and only those.

        A payload that changed under a sample that did not is a stale claim
        with a commit behind it — and a curator's own headings in the same
        body are not the Gardener's to overwrite, so this replaces section
        by section instead of replacing the body.
        """
        fresh = dataset_map(tables, samples, counts)
        manual = extract_section(fresh, MANUAL_SECTION) or ""
        sample = extract_section(fresh, SAMPLE_SECTION) or ""
        for header, section in ((MANUAL_SECTION, manual), (SAMPLE_SECTION, sample)):
            content = "\n".join(section.splitlines()[1:]).strip()
            updated = replace_section(body, header, content)
            body = updated if updated is not None else append_section(
                body, header, content)
        return body

    def _apply_content_policy(self, draft: dict, title: str,
                              is_text_source: bool) -> dict:
        policy = self.config.get("content", "inline")
        if policy == "reference" and not is_text_source:
            policy = "cached"  # converted bodies must live SOMEWHERE local
        if policy == "cached":
            self._write_body_cache(draft["id"], draft["body"])
            draft["body"] = f"# {title}"
            draft["content"] = "cached"
        elif policy == "reference":
            draft["body"] = f"# {title}"
            draft["content"] = "reference"
        return draft

    def _write_body_cache(self, node_id: str, body: str) -> None:
        if self.dry_run:
            return
        p = self.forest.body_cache_path(node_id)
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(body, encoding="utf-8", newline="\n")

    def _node_id(self, rel: Path, dest: str | None) -> str:
        parts = [slugify(p) for p in rel.parent.parts]
        prefix = [] if not dest else [dest]
        return "/".join(prefix + parts + [slugify(rel.stem)])

    @staticmethod
    def _dataset_summary(conversion: Conversion) -> str:
        """G.2.3 rule 5: the scent names the tables, not just the first one.

        A twelve-table database summarised as "table X with 6 rows" is a
        scent for the wrong thing — `locate` reads this and nothing else
        about the node.
        """
        if conversion.kind == "payload":
            tables = conversion.tables or {}
            counts = conversion.counts or {}
        else:
            tables = {t: spec["columns"] for t, spec in (conversion.schema or {}).items()}
            counts = {t: len(rows) for t, rows in (conversion.rows or {}).items()}
        named = [f"{t} ({rows_label(counts.get(t, 0))})" for t in list(tables)[:4]]
        more = len(tables) - len(named)
        listing = ", ".join(named) + (f", +{more} more" if more > 0 else "")
        noun = "table" if len(tables) == 1 else f"{len(tables)} tables"
        summary = (f"Tabular data '{conversion.title}': {noun} "
                   f"{listing}. Adopted from source; pending curation.")
        try:
            validate_summary(summary)
        except VineError:
            summary = (f"Tabular data '{conversion.title}' with {len(tables)} "
                       f"table(s). Adopted from source; pending curation.")
        return summary

    # -- sync (G.3) ----------------------------------------------------------

    def _passports(self) -> dict[str, dict]:
        """source_path -> {id, hash, size, mtime}, read from the forest itself
        (the forest IS the sync state — no side bookkeeping to drift)."""
        out: dict[str, dict] = {}
        for nid in self.forest.iter_ids():
            try:
                node = self.forest.read(nid)
            except VineError:
                continue
            fm = node.frontmatter
            if fm.get("source_path"):
                out[str(fm["source_path"])] = {
                    "id": nid,
                    "hash": str(fm.get("source_hash", "")),
                    "size": fm.get("source_size"),
                    "mtime": fm.get("source_mtime"),
                }
        return out

    def sync(self, source: str | Path | None = None,
             path: str | None = None, dest: str | None = None) -> dict:
        """`dest` overrides the adopted root's destination for files sync
        meets for the FIRST time. Files that already have a passport keep
        the branch they were planted in — sync refreshes content, it never
        moves nodes. Without the override a caller who says where a new
        document goes is silently overruled by whatever the last adopt
        recorded."""
        return _drain(self.sync_iter(source=source, path=path, dest=dest))

    def sync_iter(self, source: str | Path | None = None,
                  path: str | None = None,
                  dest: str | None = None) -> IngestSteps:
        """One document per step (G.10); `sync` is exactly "drain this".

        Resolution and containment are eager — a bad source or an escaping
        `path` fails here, before any step — while the passport read joins
        the first step: it reads the forest, and construction touches only
        the filesystem.
        """
        src = self._resolve_source(source, recorded=True)
        dest = dest or self.config.get("dest")
        report = IngestReport()

        if path:  # G.8 targeted sync: one file, the event-trigger building block
            # `path` is source-root-relative and MUST stay there. `relative_to`
            # is lexical, so `../../x` survives the join and comes back out as
            # a "relative" path — the file would be read, slugified into a
            # `node/` branch and planted. Resolving first is what makes the
            # containment real: it collapses `..` and follows symlinks.
            f = (src / Path(path)).resolve()
            if Path(path).is_absolute() or not f.is_relative_to(src):
                raise VineError(
                    E_SCHEMA, f"sync path leaves the source root: {path}",
                    hint="A targeted sync names a path relative to the "
                         "adopted source root.")
            rel = f.relative_to(src).as_posix()

            def one():
                passports = self._passports()
                before = _counts(report)
                if f.is_file():
                    self._sync_one(src, f, rel, passports, dest, report)
                elif rel in passports:
                    report.stale.append(passports[rel]["id"])
                else:
                    report.unsupported.append(rel)
                yield _step(rel, 1, 1, report, before)
                return report.as_dict()

            return IngestSteps(1, one(), report)

        files = self._walk(src)

        def steps():
            passports = self._passports()
            seen: set[str] = set()
            for i, f in enumerate(files):
                rel = f.relative_to(src).as_posix()
                seen.add(rel)
                before = _counts(report)
                self._sync_one(src, f, rel, passports, dest, report)
                yield _step(rel, i + 1, len(files), report, before)
            for rel, info in passports.items():
                if rel not in seen:
                    report.stale.append(info["id"])
            return report.as_dict()

        return IngestSteps(len(files), steps(), report)

    def _sync_one(self, src: Path, f: Path, rel: str, passports: dict,
                  dest: str | None, report: IngestReport) -> None:
        if rel not in passports:
            self._ingest_file(src, f, dest, report)
            return
        info = passports[rel]
        st = f.stat()
        # G.8 fast-path (rsync's trick): same size + mtime -> skip hashing
        if (info.get("size") == st.st_size
                and info.get("mtime") == round(st.st_mtime, 3)):
            report.unchanged.append(info["id"])
            return
        new_hash = hashlib.sha256(f.read_bytes()).hexdigest()
        if new_hash == info["hash"]:
            report.unchanged.append(info["id"])
            return
        self._update_passport(info["id"], f, new_hash, report, rel)

    def _update_passport(self, node_id: str, f: Path, new_hash: str,
                         report: IngestReport, rel: str | None = None) -> None:
        """G.3: refresh body + source_hash via the audited write path —
        curated frontmatter (summary, tags, links, confidence) is preserved.

        No `curate` stage here, and that is the contract rather than an
        omission: a refresh keeps the scent somebody already approved.
        """
        ext = f.suffix.lower()
        claimants = [c for c in self.converters if ext in c.extensions]
        if not claimants:
            report.unsupported.append(node_id)
            return
        self._stage(rel or node_id, STAGE_CONVERT)
        # G.5.1: the fallback chain covers refreshes too — a describer whose
        # endpoint is down during a sync falls back to the stub exactly as it
        # does during adopt. A refresh that errored where an adopt would have
        # planted would make the same file's fate depend on which verb found
        # it.
        conversion = None
        for conv_obj in claimants:
            terminal = conv_obj is claimants[-1]
            try:
                conversion = conv_obj.convert(f)
                break
            except VineError as e:
                if terminal:
                    report.errors.append(f"{node_id}: {e.message}")
                    return
                report.errors.append(
                    f"{node_id}: {type(conv_obj).__name__} failed, "
                    f"falling back: {e.message}")
            except Exception as e:
                if terminal:
                    report.errors.append(f"{node_id}: converter error: {e}")
                    return
                report.errors.append(
                    f"{node_id}: {type(conv_obj).__name__} failed, "
                    f"falling back: {e}")

        # J.8 (v0.48): a refresh rebuilds the body from the converter, so
        # the address is stamped again here — the same line the adopt path
        # appended, through the same helper, or a re-uploaded screenshot
        # would lose where it came from on its first refresh.
        if conversion.kind == "markdown" and rel is not None:
            conversion.markdown = self._with_provenance(conversion.markdown, rel)

        if self.dry_run:
            # An update refreshes the body and keeps the curated scent (G.3),
            # so there is no passport to review — reporting that it *would*
            # be refreshed is the whole preview.
            report.updated.append(node_id)
            return

        self._stage(rel or node_id, STAGE_PLANT)
        node = self.forest.read(node_id)
        fm = dict(node.frontmatter)
        fm["source_hash"] = new_hash
        st = f.stat()
        fm["source_size"] = st.st_size
        fm["source_mtime"] = round(st.st_mtime, 3)
        fm["updated"] = dt.date.today().isoformat()
        body = node.body

        if conversion.kind == "dataset":
            db = self.forest.payload_path(node)
            db.unlink(missing_ok=True)  # rebuild whole (sync is a bulk load)
            conn = sqlite3.connect(db)
            try:
                from monkeyllm.models import TableSchema, dataset_ddl

                schema = {t: TableSchema.model_validate(s)
                          for t, s in conversion.schema.items()}
                for stmt in dataset_ddl(schema):
                    conn.execute(stmt)
                for tname, table_rows in (conversion.rows or {}).items():
                    if table_rows:
                        ph = ", ".join("?" * len(schema[tname].columns))
                        conn.executemany(f"INSERT INTO {tname} VALUES ({ph})",
                                         [tuple(r) for r in table_rows])
                conn.commit()
            finally:
                conn.close()
            fm["payload_hash"] = hashlib.sha256(db.read_bytes()).hexdigest()
            body = self._refresh_map(
                body,
                {t: {c: ty.upper() for c, ty in s["columns"].items()}
                 for t, s in conversion.schema.items()},
                conversion.rows,
                {t: len(r) for t, r in (conversion.rows or {}).items()})
        elif conversion.kind == "payload":
            # G.2.2 rule 5: the source replaces the payload whole.
            db = self.forest.payload_path(node)
            db.parent.mkdir(parents=True, exist_ok=True)
            shutil.copyfile(f, db)
            fm["payload_hash"] = new_hash
            body = self._refresh_map(body, conversion.tables or {},
                                     conversion.samples, conversion.counts)
        elif fm.get("content") == "cached":
            self._write_body_cache(node_id, conversion.markdown)  # body stays a stub
        elif fm.get("content") == "reference":
            pass  # the body IS the source; hash bookkeeping above is enough
        else:
            body = conversion.markdown

        # G.5.1: a refreshed media file re-archives under the same condition
        # the adopt path used, or the map and the served bytes disagree — a
        # re-uploaded screenshot would keep serving the OLD image from
        # `_assets/` under a payload_hash that still validates, while the
        # new bytes sit only in disposable staging. The digest names the
        # archived copy, so a changed original lands under a new name and
        # the stale one is removed rather than left to accumulate.
        ext = f.suffix.lower()
        is_text_source = ext in MarkdownConverter.extensions
        staged_media = (PAYLOAD_TYPE_BY_EXT.get(ext) in ("image", "audio")
                        and f.resolve().is_relative_to(
                            self.forest.root.resolve() / "_derived"))
        if (not is_text_source and conversion.kind not in ("payload", "dataset")
                and (staged_media
                     or self.config.get("archive", "never") == "always")):
            old = fm.get("payload")
            payload, ptype, phash = self._archive(f, node_id)
            if ptype:
                fm.update({"payload": payload, "payload_type": ptype,
                           "payload_hash": phash})
                if old and old != payload and old.startswith(f"{ASSETS_DIR}/"):
                    stale = self.forest.path_for(node_id).parent / old
                    try:
                        stale.unlink(missing_ok=True)
                    except OSError:
                        pass  # a lingering copy is untidy, not incorrect

        assert node.path is not None
        node.path.write_text(serialize_node(fm, body), encoding="utf-8", newline="\n")
        self.vine.git.commit([node.path], f"gardener(sync): {node_id}")
        self.vine.catalog.upsert_node(self.forest.read(node_id))
        self.vine.catalog.mark_stale(node_id)
        report.updated.append(node_id)
