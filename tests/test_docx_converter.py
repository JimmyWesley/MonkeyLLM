"""G.2.1 (spec v0.12): the DOCX built-in converter."""

import pytest

pytest.importorskip("docx", reason="python-docx (ingest extra) not installed")

from docx import Document  # noqa: E402
from docx.oxml import parse_xml  # noqa: E402

from monkeyllm.gardener import (  # noqa: E402
    CommandConverter,
    DocxConverter,
    Gardener,
    builtin_converters,
    discover_converters,
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
# a floating text box (legacy VML form) — invisible to naive paragraph.text
TEXTBOX_XML = (
    f'<w:p xmlns:w="{W_NS}" xmlns:v="urn:schemas-microsoft-com:vml">'
    "<w:r><w:pict><v:shape><v:textbox><w:txbxContent>"
    "<w:p><w:r><w:t>Callout: budget frozen until June.</w:t></w:r></w:p>"
    "</w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>"
)


def build_docx(path):
    doc = Document()
    doc.add_heading("Quarterly Report", level=1)
    p = doc.add_paragraph()
    # runs fragmented mid-word — the classic PDF->DOCX conversion artifact
    p.add_run("Tou")
    p.add_run("can")
    p.add_run(" Robotics grew 18% in Q1.")
    doc.add_heading("Risks", level=2)
    doc.add_paragraph("Supply chain remains the main risk.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "region"
    table.cell(0, 1).text = "revenue"
    table.cell(1, 0).text = "north"
    table.cell(1, 1).text = "1200"
    doc.element.body.append(parse_xml(TEXTBOX_XML))
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = "LETTERHEAD PAGE HEADER"
    doc.save(str(path))


class TestDocxConverter:
    @pytest.fixture()
    def converted(self, tmp_path):
        f = tmp_path / "quarterly-report.docx"
        build_docx(f)
        return DocxConverter().convert(f)

    def test_title_from_first_heading(self, converted):
        assert converted.kind == "markdown"
        assert converted.title == "Quarterly Report"

    def test_headings_start_below_node_title(self, converted):
        assert converted.markdown.startswith("# Quarterly Report\n")
        assert "## Quarterly Report" in converted.markdown
        assert "### Risks" in converted.markdown

    def test_fragmented_runs_join_whole(self, converted):
        assert "Toucan Robotics grew 18% in Q1." in converted.markdown

    def test_table_becomes_pipe_table(self, converted):
        assert "| region | revenue |" in converted.markdown
        assert "| north | 1200 |" in converted.markdown

    def test_text_box_content_captured(self, converted):
        assert "Callout: budget frozen until June." in converted.markdown

    def test_header_excluded(self, converted):
        assert "LETTERHEAD" not in converted.markdown

    def test_empty_document(self, tmp_path):
        f = tmp_path / "empty-doc.docx"
        Document().save(str(f))
        conv = DocxConverter().convert(f)
        assert conv.title == "empty doc"
        assert conv.markdown == "# empty doc\n"

    def test_registered_as_builtin(self):
        assert any(isinstance(c, DocxConverter) for c in builtin_converters())

    def test_command_hook_outranks_builtin(self):
        convs = discover_converters(
            {"converters": {".docx": 'sometool "{input}" -o "{output}"'}})
        first = next(c for c in convs if ".docx" in c.extensions)
        assert isinstance(first, CommandConverter)


class TestDocxAdopt:
    def test_adopt_end_to_end(self, tmp_path):
        from monkeyllm.forest import init_forest
        from monkeyllm.vine import Vine

        src = tmp_path / "dump"
        src.mkdir()
        build_docx(src / "quarterly-report.docx")
        root = tmp_path / "forest"
        init_forest(root, title="F")
        vine = Vine(root, writable=True)
        try:
            report = Gardener(vine, hooks=[]).adopt(src)
            assert report["planted"] == ["quarterly-report"]
            assert not report["errors"] and not report["unsupported"]
            node = vine.forest.read("quarterly-report")
            assert node.frontmatter["type"] == "document"
            assert node.frontmatter["title"] == "Quarterly Report"
            assert "Toucan Robotics" in node.body
        finally:
            vine.close()
